"""
Document processing service - handles file parsing and text extraction
"""
import os
import mimetypes
from typing import Tuple, Optional
from pathlib import Path
import PyPDF2
import docx
from config import settings


class DocumentProcessor:
    """Process various document formats and extract text"""
    
    def __init__(self):
        self.supported_formats = settings.allowed_extensions
    
    def process_file(self, file_path: str) -> Tuple[str, str, Optional[str]]:
        """
        Process a file and extract text content
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (full_text, mime_type, error_message)
        """
        try:
            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Get file extension
            ext = Path(file_path).suffix.lower().lstrip('.')
            
            if ext not in self.supported_formats:
                return "", mime_type, f"Unsupported file format: {ext}"
            
            # Route to appropriate parser
            if ext == 'pdf':
                text = self._extract_from_pdf(file_path)
            elif ext in ['docx', 'doc']:
                text = self._extract_from_docx(file_path)
            elif ext == 'txt':
                text = self._extract_from_txt(file_path)
            else:
                return "", mime_type, f"No parser available for: {ext}"
            
            # Clean and validate text
            text = self._clean_text(text)
            
            if not text or len(text.strip()) < 10:
                return "", mime_type, "Extracted text is too short or empty"
            
            return text, mime_type, None
            
        except Exception as e:
            return "", None, f"Error processing file: {str(e)}"
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    print(f"Warning: Failed to extract page {page_num}: {e}")
                    continue
        
        return "\n\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Unable to decode text file with common encodings")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join with single newline
        cleaned = '\n'.join(lines)
        
        # Remove multiple spaces
        import re
        cleaned = re.sub(r' +', ' ', cleaned)
        
        return cleaned.strip()
    
    def detect_document_type(self, text: str, filename: str) -> str:
        """
        Detect document type based on content and filename
        
        Returns:
            Document type: contract, sop, official_document, report, other
        """
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Simple keyword-based detection
        if any(keyword in text_lower for keyword in ['contract', '合約', '協議', 'agreement']):
            return 'contract'
        elif any(keyword in text_lower for keyword in ['sop', 'standard operating procedure', '標準作業程序']):
            return 'sop'
        elif any(keyword in text_lower for keyword in ['official', '公文', 'memorandum', '函']):
            return 'official_document'
        elif any(keyword in text_lower for keyword in ['report', '報告', 'analysis', '分析']):
            return 'report'
        elif any(keyword in filename_lower for keyword in ['contract', 'sop', 'report']):
            # Fallback to filename
            if 'contract' in filename_lower:
                return 'contract'
            elif 'sop' in filename_lower:
                return 'sop'
            elif 'report' in filename_lower:
                return 'report'
        
        return 'other'
    
    def validate_file(self, file_path: str, max_size_bytes: Optional[int] = None) -> Tuple[bool, Optional[str]]:
        """
        Validate file before processing
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        if not os.path.isfile(file_path):
            return False, "Path is not a file"
        
        # Check file size
        file_size = os.path.getsize(file_path)
        max_size = max_size_bytes or settings.max_file_size_bytes
        
        if file_size > max_size:
            return False, f"File size ({file_size} bytes) exceeds maximum ({max_size} bytes)"
        
        if file_size == 0:
            return False, "File is empty"
        
        # Check extension
        ext = Path(file_path).suffix.lower().lstrip('.')
        if ext not in self.supported_formats:
            return False, f"Unsupported file format: {ext}"
        
        return True, None
